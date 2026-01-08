from flask import Flask, render_template, request, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
from docx import Document
from xhtml2pdf import pisa

app = Flask(__name__)

# Configuration base SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///depenses.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Modèle de dépense
class Depense(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    institution = db.Column(db.String(100))
    recu = db.Column(db.String(50))
    quantite = db.Column(db.String(50))
    motif = db.Column(db.String(200))
    date = db.Column(db.String(20))
    somme = db.Column(db.String(200))

with app.app_context():
    db.create_all()

# --- Fonction calcul total ---
def calcul_total(rows_list):
    total = 0
    for r in rows_list:
        parts = r.somme.replace("FC", "").replace(",", "").split()
        for p in parts:
            try:
                total += int(p)
            except:
                pass
    return total

# --- Routes principales ---
@app.route("/", methods=["GET"])
def index():
    institution = request.args.get("institution")
    date = request.args.get("date")

    query = Depense.query
    if institution:
        query = query.filter(Depense.institution.contains(institution))
    if date:
        query = query.filter(Depense.date == date)

    rows = query.all()
    total_fc = calcul_total(rows)
    return render_template("index.html", rows=rows, total_fc=total_fc)

@app.route("/add", methods=["POST"])
def add_row():
    depense = Depense(
        institution=request.form.get("institution"),
        recu=request.form.get("recu"),
        quantite=request.form.get("quantite"),
        motif=request.form.get("motif"),
        date=request.form.get("date"),
        somme=request.form.get("somme")
    )
    db.session.add(depense)
    db.session.commit()
    return redirect(url_for("index"))

# --- Supprimer une ligne ---
@app.route("/delete/<int:id>", methods=["POST"])
def delete_row(id):
    depense = Depense.query.get_or_404(id)
    db.session.delete(depense)
    db.session.commit()
    return redirect(url_for("index"))

# --- Modifier une ligne ---
@app.route("/edit/<int:id>", methods=["GET", "POST"])
def edit_row(id):
    depense = Depense.query.get_or_404(id)
    if request.method == "POST":
        depense.institution = request.form.get("institution")
        depense.recu = request.form.get("recu")
        depense.quantite = request.form.get("quantite")
        depense.motif = request.form.get("motif")
        depense.date = request.form.get("date")
        depense.somme = request.form.get("somme")
        db.session.commit()
        return redirect(url_for("index"))
    return render_template("edit.html", depense=depense)

# --- Export Word ---
def build_docx(rows_list):
    doc = Document()
    doc.add_heading("Liste des dépenses", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    headers = ["Institution", "Reçu", "Quantité", "Motif", "Date", "Somme payé"]
    for i, h in enumerate(headers):
        hdr[i].text = h

    for r in rows_list:
        cells = table.add_row().cells
        cells[0].text = r.institution
        cells[1].text = r.recu
        cells[2].text = r.quantite
        cells[3].text = r.motif
        cells[4].text = r.date
        cells[5].text = r.somme

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@app.route("/export/docx")
def export_docx():
    rows = Depense.query.all()
    bio = build_docx(rows)
    return send_file(
        bio,
        as_attachment=True,
        download_name="liste_depenses.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- Export PDF ---
def html_to_pdf(source_html):
    pdf_bytes = BytesIO()
    pisa_status = pisa.CreatePDF(src=source_html, dest=pdf_bytes)
    if pisa_status.err:
        return None
    pdf_bytes.seek(0)
    return pdf_bytes

@app.route("/export/pdf")
def export_pdf():
    rows = Depense.query.all()
    rendered = render_template("pdf_template.html", rows=rows)
    pdf_io = html_to_pdf(rendered)
    if pdf_io is None:
        return "Erreur PDF", 500
    return send_file(
        pdf_io,
        as_attachment=True,
        download_name="liste_depenses.pdf",
        mimetype="application/pdf"
    )

if __name__ == "__main__":
    app.run(debug=True)